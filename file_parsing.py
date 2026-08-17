"""Turn an uploaded CV / job-posting file into plain text.

PDF and DOCX are parsed locally (fast, free, deterministic). Screenshots
and other images are routed through Groq's vision model, since that's the
one case local libraries can't handle — there's no reliable text layer to
pull from a picture.
"""

from __future__ import annotations

import base64
import io
import mimetypes
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from llm import get_llm

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}
TEXT_EXTENSIONS = {".txt", ".md"}

# Groq's currently vision-capable model (console.groq.com/docs/vision).
# It's flagged preview/not-production-ready in Groq's own docs as of this
# writing — re-check console.groq.com/docs/models before assuming this is
# still the best (or only) option.
VISION_MODEL = "qwen/qwen3.6-27b"

OCR_PROMPT = (
    "Transcribe ALL text visible in this image verbatim, preserving line "
    "breaks and structure as best you can. This is a job posting or a CV "
    "screenshot. Do not summarize, do not add commentary or explanation — "
    "output only the transcribed text."
)


class UnsupportedFileType(ValueError):
    pass


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch to the right extractor based on file extension."""
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(file_bytes)
    if suffix == ".docx":
        return _extract_docx(file_bytes)
    if suffix in IMAGE_EXTENSIONS:
        return _extract_image(file_bytes, filename)
    if suffix in TEXT_EXTENSIONS:
        return file_bytes.decode("utf-8", errors="ignore")

    raise UnsupportedFileType(
        f"Unsupported file type '{suffix or '(none)'}'. Supported: "
        f".pdf, .docx, .txt, .md, or an image ({', '.join(sorted(IMAGE_EXTENSIONS))})."
    )


def _extract_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n\n".join(pages).strip()
    if not text:
        raise ValueError(
            "Couldn't extract any text from this PDF — it may be a scanned "
            "image with no text layer. Try uploading it as an image instead."
        )
    return text


def _extract_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Word docs often put content in tables (e.g. skills grids) — grab those too.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    text = "\n".join(paragraphs).strip()
    if not text:
        raise ValueError("Couldn't extract any text from this .docx file.")
    return text


MAX_IMAGE_DIMENSION = 2200  # px, plenty for OCR legibility
MAX_IMAGE_BYTES = 4 * 1024 * 1024  # downscale above this to stay well under Groq's 20MB limit


def _downscale_if_needed(file_bytes: bytes, mime: str) -> tuple[bytes, str]:
    """Shrink oversized screenshots so they upload fast and stay under limits.

    Re-encodes to JPEG when downscaling, so the returned mime type must be
    used as-is (don't fall back to the original filename's extension).
    """
    if len(file_bytes) <= MAX_IMAGE_BYTES:
        return file_bytes, mime

    from PIL import Image

    image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
    image.thumbnail((MAX_IMAGE_DIMENSION, MAX_IMAGE_DIMENSION))
    buf = io.BytesIO()
    image.save(buf, format="JPEG", quality=85)
    return buf.getvalue(), "image/jpeg"


def _extract_image(file_bytes: bytes, filename: str) -> str:
    mime, _ = mimetypes.guess_type(filename)
    mime = mime or "image/png"
    file_bytes, mime = _downscale_if_needed(file_bytes, mime)
    b64 = base64.b64encode(file_bytes).decode("utf-8")
    data_url = f"data:{mime};base64,{b64}"

    llm = get_llm(temperature=0, model=VISION_MODEL)
    message = {
        "role": "user",
        "content": [
            {"type": "text", "text": OCR_PROMPT},
            {"type": "image_url", "image_url": {"url": data_url}},
        ],
    }
    response = llm.invoke([message])
    text = (response.content or "").strip()
    if not text:
        raise ValueError("The vision model returned no text for this image.")
    return text
