"""Tests for local (non-network) file parsing paths. Vision/OCR extraction
isn't covered here since it requires a live Groq call."""

from __future__ import annotations

import io

import pytest
from docx import Document
from pypdf import PdfWriter

import file_parsing


def test_extract_docx_including_tables():
    doc = Document()
    doc.add_paragraph("Ada Yilmaz")
    doc.add_paragraph("Backend Software Engineer")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Python"
    table.rows[0].cells[1].text = "FastAPI"
    buf = io.BytesIO()
    doc.save(buf)

    text = file_parsing.extract_text(buf.getvalue(), "cv.docx")

    assert "Ada Yilmaz" in text
    assert "Python | FastAPI" in text


def test_extract_pdf_with_no_text_layer_raises():
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)

    with pytest.raises(ValueError):
        file_parsing.extract_text(buf.getvalue(), "blank.pdf")


def test_extract_plain_text_file():
    text = file_parsing.extract_text(b"hello world", "notes.txt")
    assert text == "hello world"


def test_unsupported_extension_raises():
    with pytest.raises(file_parsing.UnsupportedFileType):
        file_parsing.extract_text(b"whatever", "resume.exe")


def test_downscale_if_needed_shrinks_large_images(monkeypatch):
    from PIL import Image

    img = Image.new("RGB", (3000, 3000), color="blue")
    buf = io.BytesIO()
    img.save(buf, format="PNG")

    monkeypatch.setattr(file_parsing, "MAX_IMAGE_BYTES", 0)  # force the downscale branch
    shrunk_bytes, mime = file_parsing._downscale_if_needed(buf.getvalue(), "image/png")

    assert mime == "image/jpeg"
    shrunk = Image.open(io.BytesIO(shrunk_bytes))
    assert max(shrunk.size) <= file_parsing.MAX_IMAGE_DIMENSION


def test_downscale_if_needed_leaves_small_images_alone():
    from PIL import Image

    img = Image.new("RGB", (50, 50), color="green")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    original_bytes = buf.getvalue()

    result_bytes, mime = file_parsing._downscale_if_needed(original_bytes, "image/png")

    assert result_bytes == original_bytes
    assert mime == "image/png"
